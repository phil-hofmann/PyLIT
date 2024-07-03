import inspect
import pypandoc
import datetime
import json
import plotly.graph_objects as go

from docx import Document
from pylit.ui.param_map import ParamMap, Param
from pylit.global_settings import ARRAY, FLOAT_DTYPE, INT_DTYPE

def load_settings(session_state):
    settings = {}
    keys = ["workspace", "force", "interactive", "sliders", "store", "docx", "view_config_json", "view_models_json", "view_report", "view_coeffs_plot", "wide_mode"]
    for key in keys:
        if key in session_state:
            settings[key] = session_state[key]
    return settings

def save_settings(filename, settings):
    with open(filename, 'w') as file:
        print(f"settings = {settings}")
        json.dump(settings, file)
    
def load_session_state(filename):
    try:
        with open(filename, 'r') as file:
            body = json.load(file)
            print(f"body = {body}")
            return body
    except FileNotFoundError:
        return {}

def is_plotly_figure(figure):
    return isinstance(figure, go.Figure)

def create_word_doc(output_file: str, title: str, content: str):
    # TODO: Remove image_paths ...
    
    # Create a new Word document
    doc = Document()

    today = datetime.datetime.now().date()

    # Add a section to the document
    section = doc.sections[0]
    section.different_first_page = True
    section.different_odd_even = True

    # Add a header to the section
    header = section.header

    # Add a paragraph to the header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_paragraph.text = str(today) + "\t" + str(title) + "\t Pylit"
    header_paragraph.style = 'Header'

    content = content.replace('\\', '\\\\')

    # Add text content
    doc.add_paragraph(content)

    # Save the document
    doc.save(output_file)


def convert_word_to_pdf(word_file: str, pdf_file: str):
    # Convert Word document to PDF
    pypandoc.convert_file(word_file, 'pdf', outputfile=pdf_file)


def check_structure(data: dict, structure: dict) -> bool:
    """ Check if the data has the same structure as the structure. """
    for key, value in structure.items():
        if key not in data:
            return False
        if isinstance(value, tuple):
            if len(value) == 2:
                if not isinstance(data[key], value[0]):
                    return False
                if value[0] == dict and isinstance(value[1], dict):
                    if not check_structure(data[key], value[1]):
                        return False
                elif (value[0] == list or value[0] == ARRAY) and isinstance(value[1], dict):
                    for item in data[key]:
                        if not check_structure(item, value[1]):
                            return False
                elif value[0] == dict:
                    if not all(isinstance(item, value[1]) for item in data[key].values()):
                        return False
                elif value[0] == list or value[0] == ARRAY:
                    if not all(isinstance(item, value[1]) for item in data[key]):
                        return False
            else:
                if not isinstance(data[key], value):
                    return False
        else:
            if not isinstance(data[key], value):
                return False
    return True

def extract_params(func: callable) -> ParamMap:
    param_list = []
    params = inspect.signature(func).parameters
    for name, param in params.items():
        if param.annotation != inspect.Parameter.empty:
            param_type = param.annotation
            if param_type == str or param_type == int or param_type == float or param_type == bool \
                    or param_type == list or param_type == dict or param_type == tuple or param_type == set \
                        or param_type == ARRAY or param_type == FLOAT_DTYPE or param_type == INT_DTYPE:
                param_default = param.default if param.default != inspect.Parameter.empty else None
                param_list.append(Param(name=name, type=param_type, default=param_default))
    return ParamMap(param_list)

if __name__ == "__main__":
    pass
